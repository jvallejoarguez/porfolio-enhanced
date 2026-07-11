from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "javier-vallejo-cv-harvard.docx"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
HORIZONTAL_MARGIN = Inches(0.65)
VERTICAL_MARGIN = Inches(0.5)
CONTENT_WIDTH = Inches(7.2)


def set_run_font(run, size: float, *, bold=False, italic=False, underline=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return run


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.widow_control = True


def add_hyperlink(paragraph, text: str, url: str, size=8.5):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), FONT)
    properties.append(fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    properties.append(color)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    properties.append(font_size)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    properties.append(underline)

    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_right_tab(paragraph):
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT
    )


def add_rule(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "324")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "324")
    indentation.set(qn("w:hanging"), "180")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_section_heading(document: Document, text: str):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=6, after=2.5)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text.upper()), 11, bold=True)
    add_rule(paragraph)


def add_bullet(document: Document, num_id: int, text: str):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=1, line=1.02)
    paragraph.paragraph_format.keep_together = True
    properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_properties.append(level)
    num_properties.append(num)
    properties.insert(0, num_properties)
    set_run_font(paragraph.add_run(text), 9.5)


def add_company(document: Document, company: str, location: str, dates: str):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=2.5, after=0)
    paragraph.paragraph_format.keep_with_next = True
    add_right_tab(paragraph)
    set_run_font(paragraph.add_run(f"{company}, {location}"), 10.25, bold=True)
    paragraph.add_run("\t")
    set_run_font(paragraph.add_run(dates), 9.15)


def add_role(document: Document, title: str, dates: str):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=0.5)
    paragraph.paragraph_format.keep_with_next = True
    add_right_tab(paragraph)
    set_run_font(paragraph.add_run(title), 9.6, italic=True)
    if dates:
        paragraph.add_run("\t")
        set_run_font(paragraph.add_run(dates), 9.4)


def add_project(
    document: Document,
    num_id: int,
    name: str,
    link_label: str,
    link_url: str,
    stack: str,
    text: str,
):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=1.5, after=0)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(name), 9.7, bold=True)
    set_run_font(paragraph.add_run(" | "), 9.4)
    add_hyperlink(paragraph, link_label, link_url, size=9.4)
    set_run_font(paragraph.add_run(f" | {stack}"), 9.4, italic=True)
    add_bullet(document, num_id, text)


def add_label_line(document: Document, label: str, text: str):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=0.5)
    set_run_font(paragraph.add_run(f"{label}: "), 9.4, bold=True)
    set_run_font(paragraph.add_run(text), 9.4)


def generate() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = VERTICAL_MARGIN
    section.right_margin = HORIZONTAL_MARGIN
    section.bottom_margin = VERTICAL_MARGIN
    section.left_margin = HORIZONTAL_MARGIN
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    properties = document.core_properties
    properties.title = "Javier Vallejo - Full Stack Developer CV"
    properties.subject = "Harvard-style computer science CV"
    properties.author = "Javier Vallejo"
    properties.keywords = "Full Stack Developer, Svelte, TypeScript, React, Web Platforms"

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # Named override: Harvard CS resume. Single-column, ATS-friendly, black and
    # white, compact Letter geometry, Times New Roman, and no decorative chrome.
    bullet_num_id = create_bullet_numbering(document)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, after=0)
    set_run_font(name.add_run("JAVIER VALLEJO"), 21, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=1)
    set_run_font(title.add_run("FULL-STACK SOFTWARE ENGINEER"), 10, bold=True)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=2)
    set_run_font(contact.add_run("La Línea de la Concepción, Cádiz, Spain | "), 8.3)
    add_hyperlink(
        contact,
        "jvallejoarguez@gmail.com",
        "mailto:jvallejoarguez@gmail.com",
        size=8.3,
    )
    set_run_font(contact.add_run(" | "), 8.3)
    add_hyperlink(contact, "jvallejo.dev", "https://www.jvallejo.dev/", size=8.3)
    set_run_font(contact.add_run(" | "), 8.3)
    add_hyperlink(
        contact, "GitHub", "https://github.com/jvallejoarguez", size=8.3
    )
    set_run_font(contact.add_run(" | "), 8.3)
    add_hyperlink(
        contact,
        "LinkedIn",
        "https://linkedin.com/in/javier-vallejo-arguez",
        size=8.3,
    )

    add_section_heading(document, "Profile")
    profile = document.add_paragraph()
    set_paragraph_spacing(profile, after=0.5, line=1.03)
    set_run_font(
        profile.add_run(
            "Performance-minded full-stack engineer who turns complex platform constraints into fast, app-like products. I combine architecture, optimization, cross-device craft, and strong design judgment, with primary ownership of DB Games Grid and its flagship portal experience."
        ),
        9.5,
    )

    add_section_heading(document, "Experience")
    add_company(document, "DigitalBeat LTD", "Gibraltar", "Apr 2024 - Present")
    add_role(document, "Full Stack Developer", "May 2025 - Present")
    for point in [
        "Took primary ownership of DB Games Grid around its 1.4-era codebase and led its 2.x evolution into a reusable Svelte 5 platform for game discovery and live casino experiences.",
        "Built the flagship Hard Rock Bet Mexico portal across casino, live casino, sportsbook, and promotions, including custom mobile navigation, pre-paint state, staged content bootstrapping, route orchestration, and authentication-aware UI.",
        "Made performance a product constraint: kept shipped code lean, eliminated repeated work, shared timers and caches, scheduled visual updates efficiently, and optimized rendering for lower-end devices.",
        "Expanded the platform with live data, search and facets, personalization, jackpots, provider navigation, deep routes, analytics, and accessible interaction.",
        "Partnered across product, design, engineering, and delivery to refine modern responsive patterns and carry the shared foundation into additional brands.",
    ]:
        add_bullet(document, bullet_num_id, point)

    add_role(document, "Web Operations Executive", "Apr 2024 - May 2025")
    for point in [
        "Developed responsive HTML, CSS, and JavaScript campaigns and brand experiences within Playtech CMS.",
        "Improved production reliability through troubleshooting, performance checks, cross-browser QA, content versioning, and release coordination.",
    ]:
        add_bullet(document, bullet_num_id, point)

    add_company(document, "The Rock Hotel Gibraltar", "Gibraltar", "Feb 2024 - Apr 2024")
    add_role(document, "Web Developer Intern", "")
    add_bullet(
        document,
        bullet_num_id,
        "Turned the hotel's physical Wall of Fame into an interactive, touchscreen-compatible web experience and adapted it across visitor displays.",
    )

    add_company(document, "Informatica CR", "Spain (Remote)", "Sep 2023 - Dec 2023")
    add_role(document, "Web Developer Intern", "")
    add_bullet(
        document,
        bullet_num_id,
        "Built a WordPress ticket and receipt workflow with customer accounts, print requests, PDF generation, history, and document retrieval.",
    )

    add_section_heading(document, "Education")
    add_company(document, "Cesur", "Spain", "2022 - 2024")
    education = document.add_paragraph()
    set_paragraph_spacing(education, after=0.5)
    set_run_font(
        education.add_run(
            "Higher Technical Diploma in Web Application Development (DAW) | Grade: 9.8/10"
        ),
        9.5,
        italic=True,
    )

    add_section_heading(document, "Selected Projects")
    add_project(
        document,
        bullet_num_id,
        "El Impostor",
        "Live product",
        "https://juegoimpostor.app/",
        "React, TypeScript, Cloudflare Workers, Durable Objects, WebSockets",
        "Built a server-authoritative social deduction game for 3-12 players with private rooms, role-safe messaging, reconnectable sessions, timers, and voting.",
    )
    add_project(
        document,
        bullet_num_id,
        "Nosotros",
        "Private demo",
        "https://www.youtube.com/shorts/c__sfVVFmIA",
        "Next.js, Hono, PostgreSQL, Drizzle, Cloudflare",
        "Built a private iOS-first PWA with shared calendars, photos, lists, mood tracking, memories, games, and a self-hosted data layer.",
    )

    add_section_heading(document, "Technical Skills")
    add_label_line(
        document,
        "Frontend platform",
        "TypeScript, JavaScript, Svelte 5, React, Next.js, Web Components, HTML, CSS",
    )
    add_label_line(
        document,
        "Performance engineering",
        "bundle and render optimization, lazy and conditional rendering, caching and indexing, frame-scheduled updates, low-end-device optimization",
    )
    add_label_line(
        document,
        "Product systems",
        "WebSockets, search and facets, personalization, routing, authentication-aware UI, analytics, accessibility",
    )
    add_label_line(
        document,
        "Backend and data",
        "Node.js, Hono, FastAPI, Python, SQL, PostgreSQL, Supabase",
    )
    add_label_line(
        document,
        "Delivery and quality",
        "Cloudflare Workers, Durable Objects, AWS, Docker, Vercel, Git, Vite, Playwright, Vitest",
    )
    add_label_line(
        document,
        "Certifications",
        "JavaScript Algorithms and Data Structures (freeCodeCamp, 2025); Responsive Web Design (freeCodeCamp, 2024)",
    )
    add_label_line(
        document,
        "Languages spoken",
        "Spanish (native or bilingual); English (full professional proficiency)",
    )

    document.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    generate()
